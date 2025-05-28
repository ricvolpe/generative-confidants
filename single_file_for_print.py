import os
import re
from docx import Document

IN_DIR = ''

def combine_files_to_word(IN_DIR, output_filename="combined_interviews.docx"):    
    doc = Document()
    
    # Get all files in the directory
    try:
        files = os.listdir(IN_DIR)
    except FileNotFoundError:
        print(f"Error: Directory '{IN_DIR}' not found.")
        return
    
    files = [f for f in files if 'interview' in f.lower() and f.endswith('.txt')]
    
    if not files:
        print("No files found in the specified directory.")
        return
    
    # Sort files to ensure consistent ordering
    files.sort()
    
    print(f"Found {len(files)} files:")
    for file in files:
        print(f"  - {file}")
    
    # Process each diary file
    for i, filename in enumerate(files):
        file_path = os.path.join(IN_DIR, filename)
        
        match = re.match(r'([^_]+)_interview\.txt', filename, re.IGNORECASE)
        
        if match:
            person_name = match.group(1)
        
        try:
            # Read the content of the file
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read().strip()
            
            # Add title
            title = f"{person_name}"
            doc.add_heading(title, level=1)
            
            # Add content
            if content:
                doc.add_paragraph(content)
            else:
                doc.add_paragraph("(No content)")
            
            print(f"Added: {title}")
            
        except Exception as e:
            print(f"Error reading file '{filename}': {e}")
            continue
    
    # Save the document
    try:
        output_path = os.path.join(IN_DIR, output_filename)
        doc.save(output_path)
        print(f"\nSuccessfully saved combined document as: {output_path}")
    except Exception as e:
        print(f"Error saving document: {e}")

combine_files_to_word(IN_DIR)
